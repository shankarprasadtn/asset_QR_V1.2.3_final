import io
import os
import sys
import base64
from pathlib import Path
from datetime import datetime, date
from functools import wraps

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_file,
    send_from_directory,
    session,
)
import qrcode
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Inches, Pt

# -----------------------------
# Paths / Config
# -----------------------------

if getattr(sys, "frozen", False):
    BASE_DIR = Path(sys.executable).resolve().parent
    TEMPLATE_DIR = Path(sys._MEIPASS) / "templates"  # type: ignore[attr-defined]
else:
    BASE_DIR = Path(__file__).resolve().parent
    TEMPLATE_DIR = BASE_DIR / "templates"

EXCEL_FILE = BASE_DIR / "asset_data.xlsx"
QR_DIR = BASE_DIR / "qr_images"
USERS_FILE = BASE_DIR / "users.json"
PRINT_SHEET_FILE = BASE_DIR / "qr_print_sheet.docx"

QR_DIR.mkdir(exist_ok=True)

app = Flask(__name__, template_folder=str(TEMPLATE_DIR))
app.secret_key = "change_this_secret_for_production"


# -----------------------------
# Helpers
# -----------------------------

def normalize_date(value):
    """Return YYYY-MM-DD or ''."""
    if not value:
        return ""
    if isinstance(value, str):
        return value.strip().split()[0]
    if isinstance(value, (datetime, date)):
        return value.strftime("%Y-%m-%d")
    return str(value).strip().split()[0]


def init_excel():
    """Create Excel with our standard headers if missing."""
    if not EXCEL_FILE.exists():
        wb = Workbook()
        ws = wb.active
        ws.title = "Assets"
        ws.append(
            [
                "Timestamp",
                "Asset Type",
                "Serial",
                "Location",
                "Purchase Date",
                "QR Content",
            ]
        )
        wb.save(EXCEL_FILE)


def save_excel(asset_type, serial, location, qr_text, purchase_date):
    """Append one row to Excel."""
    init_excel()
    wb = load_workbook(EXCEL_FILE)
    ws = wb.active
    ws.append(
        [
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            asset_type,
            serial,
            location,
            normalize_date(purchase_date),
            qr_text,
        ]
    )
    wb.save(EXCEL_FILE)


def read_all_rows():
    """Read all rows from Excel as list of dicts."""
    if not EXCEL_FILE.exists():
        return []

    wb = load_workbook(EXCEL_FILE)
    ws = wb.active

    headers = [c.value for c in ws[1]]
    result = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        d = dict(zip(headers, row))
        result.append(
            {
                "timestamp": d.get("Timestamp"),
                "asset_type": d.get("Asset Type"),
                "serial": d.get("Serial"),
                "location": d.get("Location"),
                "purchase_date": d.get("Purchase Date"),
                "qr_content": d.get("QR Content"),
            }
        )
    return result


def generate_qr(data, serial):
    """Generate QR PNG + data URI; saved as qr_<serial>.png."""
    qr = qrcode.QRCode(version=1, box_size=10, border=3)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    data_uri = f"data:image/png;base64,{b64}"

    filename = f"qr_{serial}.png"
    img.save(QR_DIR / filename)

    return data_uri, filename


def build_print_sheet():
    """Create qr_print_sheet.docx from Excel rows + existing QR images."""
    rows = read_all_rows()
    if not rows:
        return None

    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)

    cols = 4
    valid_items = []

    # Only include rows where QR image exists
    for r in rows:
        serial = r["serial"]
        if not serial:
            continue
        qr_file = QR_DIR / f"qr_{serial}.png"
        if qr_file.exists():
            valid_items.append((r, qr_file))

    if not valid_items:
        return None

    rows_needed = (len(valid_items) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.style = "Table Grid"
    table.autofit = False

    col_width = Inches(1.75)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    r_idx = 0
    c_idx = 0

    for r, qr_path in valid_items:
        cell = table.rows[r_idx].cells[c_idx]
        serial = r["serial"]
        asset_type = r["asset_type"] or ""
        location = r["location"] or "IN110 ITC CHENNAI"
        purchase_date = normalize_date(r["purchase_date"])
        rrdd = "3204"

        # QR image
        p = cell.paragraphs[0]
        run = p.add_run()
        run.add_picture(str(qr_path), width=Inches(1.3))
        p.alignment = 1

        # Text
        lines = [
            f"Asset Name: {asset_type}",
            f"Serial Number: {serial}",
            f"Location Code: {location}",
            f"RRDD: {rrdd}",
        ]
        if purchase_date:
            lines.append(f"Purchase Date: {purchase_date}")

        for line in lines:
            p = cell.add_paragraph(line)
            p.alignment = 1
            p.runs[0].font.size = Pt(8)

        c_idx += 1
        if c_idx == cols:
            c_idx = 0
            r_idx += 1

    doc.save(PRINT_SHEET_FILE)
    return PRINT_SHEET_FILE


# -----------------------------
# Auth Helpers (simple)
# -----------------------------

def load_users():
    if not USERS_FILE.exists():
        users = {
            "admin": {"password": "Admin@123", "role": "admin"},
            "user": {"password": "User@123", "role": "user"},
        }
        USERS_FILE.write_text(json.dumps(users, indent=2), encoding="utf-8")
        return users
    import json
    try:
        return json.loads(USERS_FILE.read_text(encoding="utf-8"))
    except Exception:
        users = {
            "admin": {"password": "Admin@123", "role": "admin"},
            "user": {"password": "User@123", "role": "user"},
        }
        USERS_FILE.write_text(json.dumps(users, indent=2), encoding="utf-8")
        return users


def current_user():
    username = session.get("username")
    if not username:
        return None
    users = load_users()
    info = users.get(username)
    if not info:
        return None
    return {"username": username, "role": info.get("role", "user")}


def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not current_user():
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return wrapper


def role_required(role):
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            user = current_user()
            if not user or user.get("role") != role:
                return "Access denied", 403
            return f(*args, **kwargs)
        return wrapper
    return decorator


@app.context_processor
def inject_user():
    return {"current_user": current_user()}


# -----------------------------
# Routes
# -----------------------------

@app.route("/")
@login_required
def home():
    return redirect(url_for("index"))


@app.route("/index", methods=["GET", "POST"])
@login_required
def index():
    qr_img = None
    message = ""
    asset_type = ""
    asset_type_other = ""
    serial = ""
    location_mode = "default"
    location_custom = ""
    purchase_date = ""

    if request.method == "POST":
        asset_type = (request.form.get("asset_type") or "").strip()
        asset_type_other = (request.form.get("asset_type_other") or "").strip()
        serial = (request.form.get("serial") or "").strip()

        # Location
        location_mode = request.form.get("location_mode", "default")
        if location_mode == "other":
            location_custom = (request.form.get("location_custom") or "").strip()
            location = location_custom
        else:
            location = "IN110 ITC CHENNAI"

        purchase_date = normalize_date(request.form.get("purchase_date", ""))

        # Handle "Other" asset type
        if asset_type == "Other" and asset_type_other:
            asset_type = asset_type_other

        if asset_type and serial and location:
            qr_text = f"{asset_type}|{serial}|{location}"
            qr_img, _ = generate_qr(qr_text, serial)
            save_excel(asset_type, serial, location, qr_text, purchase_date)
            message = "QR generated and saved."
        else:
            message = "Please fill all required fields."

    return render_template(
        "index.html",
        qr_img=qr_img,
        message=message,
        asset_type=asset_type,
        asset_type_other=asset_type_other,
        serial=serial,
        location_mode=location_mode,
        location_custom=location_custom,
        purchase_date=purchase_date,
    )


@app.route("/labels")
@login_required
def labels():
    """Display label cards for all rows that have QR images."""
    rows = read_all_rows()
    items = []
    for r in rows:
        serial = r["serial"]
        if not serial:
            continue
        filename = f"qr_{serial}.png"
        if not (QR_DIR / filename).exists():
            continue

        items.append(
            {
                "serial": serial,
                "asset_type": r["asset_type"] or "",
                "location": r["location"] or "IN110 ITC CHENNAI",
                "purchase_date": normalize_date(r["purchase_date"]),
                "rrdd": "3204",
                "filename": filename,
            }
        )

    return render_template("labels.html", items=items)


@app.route("/assets")
@login_required
def assets():
    rows = read_all_rows()
    return render_template("assets.html", items=rows)


@app.route("/bulk-upload", methods=["GET", "POST"])
@login_required
@role_required("admin")
def bulk_upload():
    """
    Excel format:
    A: Serial
    B: Purchase Date
    C: Asset Type
    D: Location
    """
    message = ""
    count = 0

    if request.method == "POST":
        file = request.files.get("file")
        if not file:
            message = "No file uploaded."
        else:
            tmp = BASE_DIR / "upload_tmp.xlsx"
            file.save(tmp)

            try:
                wb = load_workbook(tmp)
                ws = wb.active
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if not row:
                        continue

                    serial = str(row[0]).strip() if row[0] else ""
                    if not serial:
                        continue

                    purchase_date = normalize_date(row[1]) if len(row) > 1 else ""
                    asset_type = str(row[2]).strip() if len(row) > 2 and row[2] else ""
                    location = (
                        str(row[3]).strip()
                        if len(row) > 3 and row[3]
                        else "IN110 ITC CHENNAI"
                    )

                    qr_text = f"{asset_type}|{serial}|{location}"
                    generate_qr(qr_text, serial)
                    save_excel(asset_type, serial, location, qr_text, purchase_date)
                    count += 1

                message = f"Successfully processed {count} rows."

            except Exception as e:
                message = f"Error processing file: {e}"

            finally:
                if tmp.exists():
                    tmp.unlink()

    return render_template("bulk_upload.html", message=message)


@app.route("/print-sheet")
@login_required
@role_required("admin")
def print_sheet():
    output = build_print_sheet()
    if not output:
        return "No QR images found (or no matching Excel rows).", 400

    return send_file(
        output,
        as_attachment=True,
        download_name="qr_print_sheet.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


@app.route("/qr/<path:filename>")
@login_required
def qr_image(filename):
    return send_from_directory(QR_DIR, filename)


# -----------------------------
# Auth routes
# -----------------------------

@app.route("/login", methods=["GET", "POST"])
def login():
    message = ""
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        password = (request.form.get("password") or "").strip()
        users = load_users()
        user = users.get(username)
        if user and user.get("password") == password:
            session["username"] = username
            return redirect(url_for("index"))
        else:
            message = "Invalid username or password."
    return render_template("login.html", message=message)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/health")
def health():
    return {"status": "ok"}


if __name__ == "__main__":
    app.run(debug=True)
